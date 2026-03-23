"""
Browser-facing admin UI router.

Serves Jinja2 HTML pages for campaign and analysis management.
Auth: JWT stored in an HttpOnly cookie (callcoach_token).

Routes:
  GET  /admin/ui/login                      — login form
  POST /admin/ui/login                      — submit credentials
  GET  /admin/ui/logout                     — clear cookie, redirect
  GET  /admin/ui/campaigns                  — list all configurations
  GET  /admin/ui/campaigns/new              — new campaign form
  POST /admin/ui/campaigns/new              — create campaign
  GET  /admin/ui/campaigns/{id}/edit        — edit form
  POST /admin/ui/campaigns/{id}/edit        — update campaign
  POST /admin/ui/campaigns/{id}/delete      — delete campaign
  GET  /admin/ui/global                     — list global documents
  GET  /admin/ui/global/new                 — new global document form
  POST /admin/ui/global/new                 — create global document
  GET  /admin/ui/global/{id}/edit           — edit global document
  POST /admin/ui/global/{id}/edit           — update global document
  POST /admin/ui/global/{id}/delete         — delete global document
  POST /admin/ui/upload-extract             — extract text from uploaded file
  GET  /admin/ui/analyses                   — list analyses (last 100)
  GET  /admin/ui/analyses/{id}              — analysis detail
"""

import asyncio
import io
import logging
import re
from typing import Optional

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from jose import JWTError, jwt
from sqlalchemy import delete as sa_delete, desc, func, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm.attributes import flag_modified

from config import settings
from database import get_db
from models import Analysis, Campaign, GlobalDocument, PromptSection
from utils.auth import create_access_token, verify_admin_password

router = APIRouter(prefix="/admin/ui", tags=["admin-ui"])
templates = Jinja2Templates(directory="templates")
logger = logging.getLogger(__name__)

# ── Label colors — used in all templates that show Acuity labels ─────────────
# These match the exact badge colors shown in the Acuity label picker UI.
# Update here to keep ALL pages in sync automatically.
LABEL_COLORS: dict[str, str] = {
    "PRESO":           "#f9a825",   # amber
    "CONFERMATO":      "#1e88e5",   # medium blue
    "APP.TO OK":       "#388e3c",   # dark green
    "APP.TO KO":       "#c62828",   # dark red
    "ANNULLATO":       "#f06292",   # hot pink
    "DA RICHIAMARE":   "#7c3aed",   # purple
    "NO SHOW":         "#6b7280",   # gray
}
# Fallback: Acuity restituisce anche un campo "color" con nomi CSS (yellow, blue, etc.)
# Usato se il nome label non combacia col dizionario sopra.
_ACUITY_CSS_COLOR_MAP: dict[str, str] = {
    "yellow":     "#f9a825",
    "orange":     "#f97316",
    "pink":       "#f06292",
    "red":        "#c62828",
    "green":      "#388e3c",
    "teal":       "#0d9488",
    "blue":       "#1e88e5",
    "purple":     "#7c3aed",
    "gray":       "#6b7280",
    "grey":       "#6b7280",
    "black":      "#374151",
    "cream":      "#d4a55a",
}
_LABEL_COLORS_DEFAULT = "#708090"   # fallback for unknown labels

_COOKIE_NAME = "callcoach_token"
_ALGORITHM = "HS256"


# ── Auth helper ───────────────────────────────────────────────────────────────

def _is_admin(request: Request) -> bool:
    """Return True if the request carries a valid admin JWT cookie."""
    token = request.cookies.get(_COOKIE_NAME)
    if not token:
        return False
    try:
        payload = jwt.decode(token, settings.secret_key, algorithms=[_ALGORITHM])
        return payload.get("role") == "admin"
    except (JWTError, Exception):
        return False


def _login_redirect() -> RedirectResponse:
    return RedirectResponse(url="/admin/ui/login", status_code=303)


# ── Login / Logout ────────────────────────────────────────────────────────────

@router.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    # Already authenticated → go straight to main page
    if _is_admin(request):
        return RedirectResponse(url="/admin/ui/main", status_code=303)
    return templates.TemplateResponse("login.html", {"request": request, "error": ""})


@router.post("/login")
async def login_submit(request: Request, password: str = Form(...)):
    if not verify_admin_password(password):
        return templates.TemplateResponse(
            "login.html",
            {"request": request, "error": "Password errata. Riprova."},
            status_code=401,
        )
    token = create_access_token({"role": "admin", "sub": "admin"})
    response = RedirectResponse(url="/admin/ui/main", status_code=303)
    response.set_cookie(
        _COOKIE_NAME,
        token,
        httponly=True,
        max_age=86_400,   # 24 hours
        samesite="lax",
    )
    return response


@router.get("/logout")
async def logout():
    response = RedirectResponse(url="/admin/ui/login", status_code=303)
    response.delete_cookie(_COOKIE_NAME)
    return response


# ── Campaign list ─────────────────────────────────────────────────────────────

@router.get("/campaigns", response_class=HTMLResponse)
async def campaigns_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Campaign).order_by(Campaign.code))
    campaigns = result.scalars().all()

    # Flash message from URL query param (set after create/update/delete)
    flash_ok = request.query_params.get("ok", "")
    flash_err = request.query_params.get("err", "")

    return templates.TemplateResponse(
        "campaigns_list.html",
        {
            "request": request,
            "campaigns": campaigns,
            "active_page": "campaigns",
            "flash_ok": flash_ok,
            "flash_err": flash_err,
        },
    )


# ── New campaign ──────────────────────────────────────────────────────────────

@router.get("/campaigns/new", response_class=HTMLResponse)
async def campaign_new_form(request: Request):
    if not _is_admin(request):
        return _login_redirect()
    return templates.TemplateResponse(
        "campaigns_form.html",
        {
            "request": request,
            "campaign": None,
            "active_page": "campaigns",
            "flash_err": None,
            "form_data": None,
        },
    )


@router.post("/campaigns/new")
async def campaign_new_submit(
    request: Request,
    db: AsyncSession = Depends(get_db),
    code: str = Form(...),
    nome: str = Form(""),
    script: str = Form(""),
    qualification_params: str = Form(""),
    client_info: str = Form(""),
    email_recipients_raw: str = Form(""),
    email_no_operator: str = Form("off"),
    email_disabled: str = Form("off"),
    notes: str = Form(""),
    prompt_extra: str = Form(""),
    transcription_engine: str = Form(""),
    active: str = Form("off"),
):
    if not _is_admin(request):
        return _login_redirect()

    code = code.strip().upper()

    # Validate
    if not code:
        return templates.TemplateResponse(
            "campaigns_form.html",
            {
                "request": request,
                "campaign": None,
                "active_page": "campaigns",
                "flash_err": "Il codice è obbligatorio.",
                "form_data": _form_snapshot(locals()),
            },
            status_code=422,
        )

    # Check duplicate
    existing = await db.execute(select(Campaign).where(Campaign.code == code))
    if existing.scalar_one_or_none():
        return templates.TemplateResponse(
            "campaigns_form.html",
            {
                "request": request,
                "campaign": None,
                "active_page": "campaigns",
                "flash_err": f"Esiste già una configurazione con codice «{code}».",
                "form_data": _form_snapshot(locals()),
            },
            status_code=422,
        )

    recipients = _parse_recipients(email_recipients_raw)
    campaign = Campaign(
        code=code,
        type=code.split("-")[0],          # derive type from first segment
        nome=nome.strip() or None,
        script=script.strip() or None,
        qualification_params=qualification_params.strip() or None,
        client_info=client_info.strip() or None,
        email_recipients=recipients or None,
        email_no_operator=(email_no_operator == "on"),
        email_disabled=(email_disabled == "on"),
        notes=notes.strip() or None,
        prompt_extra=prompt_extra.strip() or None,
        transcription_engine=transcription_engine.strip() or None,
        active=(active == "on"),
    )
    db.add(campaign)
    await db.commit()

    from urllib.parse import quote
    msg = quote(f"Configurazione «{code}» creata con successo.")
    return RedirectResponse(url=f"/admin/ui/campaigns?ok={msg}", status_code=303)


# ── Edit campaign ─────────────────────────────────────────────────────────────

@router.get("/campaigns/{campaign_id}/edit", response_class=HTMLResponse)
async def campaign_edit_form(
    campaign_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Campaign).where(Campaign.id == campaign_id))
    campaign = result.scalar_one_or_none()
    if not campaign:
        return RedirectResponse(url="/admin/ui/campaigns", status_code=303)

    return templates.TemplateResponse(
        "campaigns_form.html",
        {
            "request": request,
            "campaign": campaign,
            "active_page": "campaigns",
            "flash_err": None,
            "form_data": None,
        },
    )


@router.post("/campaigns/{campaign_id}/edit")
async def campaign_edit_submit(
    campaign_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    nome: str = Form(""),
    script: str = Form(""),
    qualification_params: str = Form(""),
    client_info: str = Form(""),
    email_recipients_raw: str = Form(""),
    email_no_operator: str = Form("off"),
    email_disabled: str = Form("off"),
    notes: str = Form(""),
    prompt_extra: str = Form(""),
    transcription_engine: str = Form(""),
    active: str = Form("off"),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Campaign).where(Campaign.id == campaign_id))
    campaign = result.scalar_one_or_none()
    if not campaign:
        return RedirectResponse(url="/admin/ui/campaigns", status_code=303)

    campaign.nome = nome.strip() or None
    campaign.script = script.strip() or None
    campaign.qualification_params = qualification_params.strip() or None
    campaign.client_info = client_info.strip() or None
    campaign.email_recipients = _parse_recipients(email_recipients_raw) or None
    campaign.email_no_operator = (email_no_operator == "on")
    campaign.email_disabled = (email_disabled == "on")
    campaign.notes = notes.strip() or None
    campaign.prompt_extra = prompt_extra.strip() or None
    campaign.transcription_engine = transcription_engine.strip() or None
    campaign.active = (active == "on")

    await db.commit()

    from urllib.parse import quote
    msg = quote(f"Configurazione «{campaign.code}» aggiornata.")
    return RedirectResponse(url=f"/admin/ui/campaigns?ok={msg}", status_code=303)


# ── Test email ────────────────────────────────────────────────────────────────

@router.get("/test-email", response_class=JSONResponse)
async def test_email(request: Request):
    if not _is_admin(request):
        return JSONResponse({"error": "non autorizzato"}, status_code=403)

    from config import settings as cfg
    from services.email_service import send_analysis_report

    to_addr = request.query_params.get("to", cfg.fallback_email)
    try:
        await send_analysis_report(
            recipients=[to_addr],
            html_content="<h1>Test CallCoach AI</h1><p>Email di test inviata correttamente.</p>",
            operator_name="Test",
            qualification_level="buona",
            appointment_datetime="2026-01-01",
        )
        return JSONResponse({
            "ok": True,
            "to": to_addr,
            "brevo_key_set": bool(cfg.brevo_api_key),
        })
    except Exception as exc:
        return JSONResponse({
            "ok": False,
            "error": str(exc),
            "brevo_key_set": bool(cfg.brevo_api_key),
        }, status_code=500)


# ── Debug: raw DB + Acuity data (remove after diagnosis) ─────────────────────

@router.get("/debug")
async def debug_data(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Temporary diagnostic: returns JSON with exact DB campaign codes and raw Acuity types."""
    if not _is_admin(request):
        return JSONResponse({"error": "auth"}, status_code=401)

    from datetime import datetime, timedelta, timezone
    from fastapi.responses import JSONResponse as _JSON
    from services.acuity import list_appointments, clear_appointments_cache

    # DB campaigns
    camp_result = await db.execute(select(Campaign))
    camps = camp_result.scalars().all()
    db_campaigns = [
        {
            "id": c.id,
            "code": c.code,
            "code_repr": repr(c.code),
            "active": c.active,
            "active_type": type(c.active).__name__,
        }
        for c in camps
    ]

    # Acuity appointments (last 7 days, first 20)
    clear_appointments_cache()
    now = datetime.now(timezone.utc)
    min_date = (now - timedelta(days=7)).strftime("%Y-%m-%d")
    max_date = (now + timedelta(days=30)).strftime("%Y-%m-%d")
    try:
        appts1 = await list_appointments(1, min_date=min_date, max_date=max_date, max_results=20)
        appts2 = await list_appointments(2, min_date=min_date, max_date=max_date, max_results=20)
        appts = appts1 + appts2
    except Exception as exc:
        appts = []

    from services.acuity import find_opr_field, get_operator_display

    acuity_appointments = []
    for a in appts[:30]:
        op_result = get_operator_display(a)
        all_form_fields = [
            {
                "form_name": form.get("name", ""),
                "field_name": v.get("name", ""),
                "value": v.get("value", ""),
            }
            for form in (a.get("forms") or [])
            for v in (form.get("values") or [])
        ]
        acuity_appointments.append({
            "id": a.get("id"),
            "type": a.get("type"),
            "find_opr_field_result": find_opr_field(a),
            "get_operator_display_result": op_result,
            "all_form_fields": all_form_fields,
        })

    return _JSON({"db_campaigns": db_campaigns, "acuity_appointments": acuity_appointments})


# ── Delete campaign ───────────────────────────────────────────────────────────

@router.post("/campaigns/{campaign_id}/delete")
async def campaign_delete(
    campaign_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Campaign).where(Campaign.id == campaign_id))
    campaign = result.scalar_one_or_none()
    if not campaign:
        return RedirectResponse(url="/admin/ui/campaigns", status_code=303)

    code = campaign.code
    await db.delete(campaign)
    await db.commit()

    from urllib.parse import quote
    msg = quote(f"Configurazione «{code}» eliminata.")
    return RedirectResponse(url=f"/admin/ui/campaigns?ok={msg}", status_code=303)


# ── Duplicate campaign ────────────────────────────────────────────────────────

@router.post("/campaigns/{campaign_id}/duplicate")
async def campaign_duplicate(
    campaign_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    new_code: str = Form(...),
):
    if not _is_admin(request):
        return _login_redirect()

    from urllib.parse import quote

    new_code = new_code.strip().upper()
    if not new_code:
        msg = quote("Il codice non può essere vuoto.")
        return RedirectResponse(url=f"/admin/ui/campaigns?err={msg}", status_code=303)

    # Load original
    result = await db.execute(select(Campaign).where(Campaign.id == campaign_id))
    original = result.scalar_one_or_none()
    if not original:
        return RedirectResponse(url="/admin/ui/campaigns", status_code=303)

    # Check new code uniqueness
    existing = await db.execute(select(Campaign).where(Campaign.code == new_code))
    if existing.scalar_one_or_none():
        msg = quote(f"Esiste già una configurazione con codice «{new_code}».")
        return RedirectResponse(url=f"/admin/ui/campaigns?err={msg}", status_code=303)

    new_camp = Campaign(
        code=new_code,
        nome=original.nome,
        type=original.type,
        script=original.script,
        qualification_params=original.qualification_params,
        client_info=original.client_info,
        email_recipients=list(original.email_recipients) if original.email_recipients else [],
        notes=original.notes,
        prompt_extra=original.prompt_extra,
        active=original.active,
    )
    db.add(new_camp)
    await db.commit()

    msg = quote(f"Configurazione «{new_code}» creata come copia di «{original.code}».")
    return RedirectResponse(url=f"/admin/ui/campaigns?ok={msg}", status_code=303)


# ── Prompt editor ─────────────────────────────────────────────────────────────

@router.get("/prompt", response_class=HTMLResponse)
async def prompt_editor(request: Request, db: AsyncSession = Depends(get_db)):
    if not _is_admin(request):
        return _login_redirect()
    from services.prompt_db import SECTION_METADATA, _DEFAULT_SECTIONS, get_prompt_sections
    current_sections = await get_prompt_sections()
    sections_for_template = [
        {
            **meta,
            "content": current_sections.get(meta["key"], _DEFAULT_SECTIONS.get(meta["key"], "")),
            "default_content": _DEFAULT_SECTIONS.get(meta["key"], ""),
        }
        for meta in SECTION_METADATA
    ]
    return templates.TemplateResponse("prompt_editor.html", {
        "request": request,
        "sections": sections_for_template,
        "active_page": "prompt",
        "flash_ok": request.query_params.get("ok", ""),
        "flash_err": request.query_params.get("err", ""),
    })


@router.post("/prompt")
async def prompt_editor_save(request: Request, db: AsyncSession = Depends(get_db)):
    if not _is_admin(request):
        return _login_redirect()
    from services.prompt_db import SECTION_METADATA, clear_prompt_sections_cache
    from models import PromptSection
    form = await request.form()
    for idx, meta in enumerate(SECTION_METADATA):
        key = meta["key"]
        content = (form.get(key) or "").strip()
        existing = await db.execute(select(PromptSection).where(PromptSection.section_key == key))
        row = existing.scalar_one_or_none()
        if row:
            row.content = content
        else:
            db.add(PromptSection(
                section_key=key,
                title=meta["title"],
                content=content,
                sort_order=idx,
            ))
    await db.commit()
    clear_prompt_sections_cache()
    from urllib.parse import quote
    return RedirectResponse(url=f"/admin/ui/prompt?ok={quote('Prompt salvato con successo.')}", status_code=303)


# ── Analyses list ─────────────────────────────────────────────────────────────

@router.get("/analyses", response_class=HTMLResponse)
async def analyses_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(
        select(Analysis).order_by(desc(Analysis.created_at)).limit(100)
    )
    analyses = result.scalars().all()

    from utils.helpers import utcnow
    return templates.TemplateResponse(
        "analyses_list.html",
        {
            "request": request,
            "analyses": analyses,
            "now": utcnow(),
            "active_page": "analyses",
            "flash_ok": request.query_params.get("ok", ""),
            "flash_err": request.query_params.get("err", ""),
        },
    )


# ── Reinvia email da lista analisi ────────────────────────────────────────────

@router.post("/analyses/{analysis_id}/send-email", response_class=JSONResponse)
async def analysis_send_email(
    analysis_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    mode: str = "operator",   # "operator" | "inoltro"
):
    if not _is_admin(request):
        return JSONResponse({"ok": False, "error": "non autorizzato"}, status_code=403)

    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if not analysis:
        return JSONResponse({"ok": False, "error": "analisi non trovata"}, status_code=404)
    if not analysis.report_html:
        return JSONResponse({"ok": False, "error": "nessun report HTML disponibile"}, status_code=400)

    from services.email_service import send_analysis_report

    _INOLTRO = "inoltro@effoncall.com"

    if mode == "inoltro":
        recipients = [_INOLTRO]
    else:
        # mode == "operator": usa operator_email salvato in fase di analisi
        op_email = (analysis.operator_email or "").strip()
        # Fallback per analisi vecchie senza operator_email: ricava da operator_name
        # operator_name formato: "45-FRANCESCA F." → op.45.francesca@effoncall.com
        if not op_email and analysis.operator_name:
            import re as _re
            m = _re.match(r'^(\d+)\s*-\s*(\S+)', analysis.operator_name.strip())
            if m:
                num       = m.group(1).zfill(2)
                firstname = m.group(2).strip(".").lower()
                op_email  = f"op.{num}.{firstname}@effoncall.com"
        if not op_email:
            return JSONResponse({"ok": False, "error": "indirizzo email operatore non ricavabile"}, status_code=400)
        recipients = [op_email]

    try:
        await send_analysis_report(
            recipients=recipients,
            html_content=analysis.report_html,
            operator_name=analysis.operator_name or "—",
            qualification_level=analysis.qualification_level or "buona",
            appointment_datetime=str(analysis.appointment_datetime or ""),
        )
        return JSONResponse({"ok": True, "recipients": recipients})
    except Exception as exc:
        return JSONResponse({"ok": False, "error": str(exc)}, status_code=500)


# ── Archivio (tutte le analisi, sola lettura) ─────────────────────────────────

@router.get("/archivio", response_class=HTMLResponse)
async def archivio(
    request: Request,
    db: AsyncSession = Depends(get_db),
    mese: str = "",   # "YYYY-MM"
    op: str = "",
    camp: str = "",
    qual: str = "",
):
    if not _is_admin(request):
        return _login_redirect()

    from datetime import datetime as _dt
    from sqlalchemy import and_, extract

    stmt = select(Analysis).order_by(desc(Analysis.created_at))

    filters = []
    if mese:
        try:
            yr, mo = int(mese[:4]), int(mese[5:7])
            filters.append(extract("year",  Analysis.created_at) == yr)
            filters.append(extract("month", Analysis.created_at) == mo)
        except (ValueError, IndexError):
            pass
    if op:
        filters.append(Analysis.operator_name.ilike(f"%{op}%"))
    if camp:
        filters.append(Analysis.campaign_code.ilike(f"%{camp}%"))
    if qual:
        filters.append(Analysis.qualification_level == qual)

    if filters:
        stmt = stmt.where(and_(*filters))

    result = await db.execute(stmt)
    analyses = result.scalars().all()

    # Mesi disponibili: ricavati in Python dagli oggetti già caricati
    # (evita query raw TO_CHAR che causa problemi di transazione su Supabase async)
    _all_months_result = await db.execute(
        select(Analysis.created_at).where(Analysis.created_at.is_not(None))
    )
    _all_dates = _all_months_result.scalars().all()
    available_months = sorted(
        set(d.strftime("%Y-%m") for d in _all_dates if d),
        reverse=True,
    )

    from utils.helpers import utcnow
    return templates.TemplateResponse(
        "archivio.html",
        {
            "request": request,
            "analyses": analyses,
            "now": utcnow(),
            "active_page": "archivio",
            "available_months": available_months,
            "sel_mese": mese,
            "sel_op": op,
            "sel_camp": camp,
            "sel_qual": qual,
        },
    )


# ── Delete single analysis ────────────────────────────────────────────────────

@router.post("/analyses/{analysis_id}/delete")
async def analysis_delete(
    analysis_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if analysis:
        await db.delete(analysis)
        await db.commit()

    from urllib.parse import quote
    msg = quote(f"Analisi #{analysis_id} eliminata.")
    return RedirectResponse(url=f"/admin/ui/analyses?ok={msg}", status_code=303)


# ── Clear all analyses ────────────────────────────────────────────────────────

@router.post("/analyses/clear")
async def clear_analyses(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    count_result = await db.execute(select(func.count()).select_from(Analysis))
    count = count_result.scalar_one_or_none() or 0
    await db.execute(sa_delete(Analysis))
    await db.commit()

    from urllib.parse import quote
    msg = quote(f"{count} analisi eliminate.")
    return RedirectResponse(url=f"/admin/ui/analyses?ok={msg}", status_code=303)


# ── Analysis detail ───────────────────────────────────────────────────────────

@router.get("/analyses/{analysis_id}", response_class=HTMLResponse)
async def analysis_detail(
    analysis_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    ok: str = "",
    err: str = "",
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if not analysis:
        return RedirectResponse(url="/admin/ui/analyses", status_code=303)

    return templates.TemplateResponse(
        "analysis_detail.html",
        {
            "request": request,
            "analysis": analysis,
            "active_page": "analyses",
            "flash_ok": ok,
            "flash_err": err,
        },
    )


# ── Mark as errore tecnico ────────────────────────────────────────────────────

@router.post("/analyses/{analysis_id}/mark-errore-tecnico")
async def mark_errore_tecnico(
    analysis_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    from urllib.parse import quote

    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if not analysis:
        return RedirectResponse(url="/admin/ui/analyses", status_code=303)

    # Flag in report_json + reset fuori_parametro so banner won't show as NON IN TARGET
    rj = dict(analysis.report_json or {})
    rj["errore_tecnico"] = True
    if "qualificazione" in rj:
        q = dict(rj["qualificazione"])
        q["fuori_parametro"] = False
        q["spiegazione"] = "Analisi non valida — errore tecnico di trascrizione."
        rj["qualificazione"] = q
    analysis.report_json = rj
    flag_modified(analysis, "report_json")
    analysis.qualification_level = "errore_tecnico"
    await db.commit()

    msg = quote("Analisi contrassegnata come errore tecnico.")
    return RedirectResponse(url=f"/admin/ui/analyses/{analysis_id}?ok={msg}", status_code=303)


# ── Re-run analysis from detail page ─────────────────────────────────────────

@router.post("/analyses/{analysis_id}/rianalizza")
async def rianalizza_from_detail(
    analysis_id: int,
    request: Request,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    from services.acuity import get_appointment
    from routers.webhook import run_analysis_pipeline
    from urllib.parse import quote

    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if not analysis:
        return RedirectResponse(url="/admin/ui/analyses", status_code=303)

    full_appointment = await get_appointment(analysis.appointment_id, analysis.acuity_account)
    if not full_appointment:
        msg = quote(f"Impossibile recuperare l'appuntamento {analysis.appointment_id} da Acuity.")
        return RedirectResponse(url=f"/admin/ui/analyses/{analysis_id}?err={msg}", status_code=303)

    background_tasks.add_task(
        run_analysis_pipeline,
        appointment_data=full_appointment,
        acuity_account=analysis.acuity_account,
    )

    msg = quote("Rianalisi avviata. Aggiorna la pagina tra qualche minuto per il risultato aggiornato.")
    return RedirectResponse(url=f"/admin/ui/analyses/{analysis_id}?ok={msg}", status_code=303)


# ── Analysis print / PDF ──────────────────────────────────────────────────────

@router.get("/analyses/{analysis_id}/print", response_class=HTMLResponse)
async def analysis_print(
    analysis_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """
    Serve the analysis report as a standalone HTML page suitable for
    browser print-to-PDF.  Transcript and raw JSON are excluded.
    """
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if not analysis or not analysis.report_html:
        return RedirectResponse(url=f"/admin/ui/analyses/{analysis_id}", status_code=303)

    # Inject print controls and print CSS into the report HTML
    report_id_label = f"Report #{analysis_id}"
    if analysis.campaign_code:
        report_id_label += f" &mdash; {analysis.campaign_code}"

    print_bar = f"""
<style>
  @media print {{
    .ec-print-bar {{ display: none !important; }}
    body {{ background: #fff !important; }}
  }}
</style>
<div class="ec-print-bar" style="
  position:sticky; top:0; z-index:100;
  background:#001126; color:#fff;
  padding:8px 16px;
  display:flex; align-items:center; justify-content:space-between;
  font-family:Arial,sans-serif; font-size:12px;
">
  <span style="color:rgba(255,255,255,.6)">{report_id_label}</span>
  <div style="display:flex;gap:8px">
    <button onclick="window.print()" style="
      background:#fff; color:#001126; border:none;
      padding:5px 14px; border-radius:4px; cursor:pointer;
      font-size:12px; font-weight:700; letter-spacing:.3px;
    ">Stampa / Salva PDF</button>
    <button onclick="window.close()" style="
      background:rgba(255,255,255,.15); color:#fff; border:none;
      padding:5px 12px; border-radius:4px; cursor:pointer; font-size:12px;
    ">Chiudi</button>
  </div>
</div>
"""

    html = analysis.report_html
    if "<body>" in html:
        html = html.replace("<body>", f"<body>\n{print_bar}", 1)
    else:
        html = print_bar + html

    return HTMLResponse(content=html)


# ── Appointments v2 shell ─────────────────────────────────────────────────────

@router.get("/appointments2", response_class=HTMLResponse)
async def appointments2_list(request: Request):
    if not _is_admin(request):
        return _login_redirect()
    return templates.TemplateResponse(
        "appointments_v2.html",
        {"request": request, "active_page": "appointments2"},
    )


# ── Appointments list (instant shell — no Acuity calls) ──────────────────────

@router.get("/appointments", response_class=HTMLResponse)
async def appointments_list(request: Request):
    """Render the shell instantly; data loads async via /appointments/data."""
    if not _is_admin(request):
        return _login_redirect()
    return templates.TemplateResponse(
        "appointments_list.html",
        {
            "request": request,
            "active_page": "appointments",
            "flash_ok": request.query_params.get("ok", ""),
            "flash_err": request.query_params.get("err", ""),
        },
    )


# ── Appointments data fragment (called async by JS) ───────────────────────────

@router.get("/appointments/data", response_class=HTMLResponse)
async def appointments_data(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Return the appointments table HTML fragment (heavy, called via fetch())."""
    if not _is_admin(request):
        return HTMLResponse(status_code=401)

    from datetime import datetime, timedelta, timezone

    from services.acuity import (
        clear_appointments_cache,
        extract_phone,
        extract_piva,
        extract_ragione_sociale,
        get_operator_display,
        list_appointments,
    )
    from services.campaign_parser import parse_campaign_code
    from utils.helpers import parse_iso_datetime

    # ?refresh=1 forces a fresh Acuity fetch by clearing the in-memory cache
    if request.query_params.get("refresh") == "1":
        clear_appointments_cache()

    now = datetime.now(timezone.utc)
    min_date = (now - timedelta(days=30)).strftime("%Y-%m-%d")
    max_date = (now + timedelta(days=365)).strftime("%Y-%m-%d")

    # Fetch from both Acuity accounts in parallel
    acuity_results = await asyncio.gather(
        list_appointments(1, min_date=min_date, max_date=max_date),
        list_appointments(2, min_date=min_date, max_date=max_date),
        return_exceptions=True,
    )
    appts_1 = acuity_results[0] if isinstance(acuity_results[0], list) else []
    appts_2 = acuity_results[1] if isinstance(acuity_results[1], list) else []

    for a in appts_1:
        a["_account"] = 1
    for a in appts_2:
        a["_account"] = 2

    all_appts = sorted(appts_1 + appts_2, key=lambda a: a.get("datetime", ""), reverse=True)

    # Load ALL campaigns (active + inactive) — one query
    camp_result = await db.execute(
        select(Campaign).where(Campaign.code != "_GLOBAL_")
    )
    _all_camps = list(camp_result.scalars().all())
    # Normalise keys: strip + uppercase so DB entries with accidental whitespace
    # or wrong casing (e.g. "AVANZ-AVI-0000 ") still match correctly.
    # Treat active=NULL as active=TRUE (matches schema DEFAULT TRUE intent).
    all_campaigns: dict[str, Campaign] = {c.code.strip().upper(): c for c in _all_camps if c.active is not False}
    all_campaigns_inactive: dict[str, Campaign] = {c.code.strip().upper(): c for c in _all_camps if c.active is False}

    logger.info(
        "appointments/data: %d active campaigns loaded: %s | raw active values: %s",
        len(all_campaigns),
        sorted(all_campaigns.keys()),
        {c.code: c.active for c in _all_camps},
    )

    # Existing analyses for these appointments — one query
    appt_ids = [str(a["id"]) for a in all_appts]
    if appt_ids:
        # Try full query first (with new columns); fall back to base columns if
        # the DB migration hasn't run yet (avoids 500 on first deploy).
        _full_cols = True
        try:
            ana_result = await db.execute(
                select(Analysis.appointment_id, Analysis.processing_status, Analysis.id,
                       Analysis.progress, Analysis.step_message, Analysis.created_at,
                       Analysis.qualification_level, Analysis.pipeline_steps,
                       Analysis.num_recordings, Analysis.total_talk_seconds,
                       Analysis.label_name, Analysis.label_color)
                .where(Analysis.appointment_id.in_(appt_ids))
            )
        except Exception as _qe:
            logger.warning("analyses full query failed (missing columns?): %s — retrying base", _qe)
            _full_cols = False
            await db.rollback()
            ana_result = await db.execute(
                select(Analysis.appointment_id, Analysis.processing_status, Analysis.id,
                       Analysis.progress, Analysis.step_message, Analysis.created_at,
                       Analysis.qualification_level)
                .where(Analysis.appointment_id.in_(appt_ids))
            )

        def _build_map_row(row) -> dict:
            return {
                "status": row.processing_status,
                "id": row.id,
                "progress": row.progress or 0,
                "step_message": row.step_message or "",
                "created_at": row.created_at,
                "created_display": row.created_at.strftime("%d/%m/%Y %H:%M") if row.created_at else "—",
                "qualification_level": row.qualification_level or "",
                "pipeline_steps": (row.pipeline_steps or {}) if _full_cols else {},
                "num_recordings": (row.num_recordings or 0) if _full_cols else 0,
                "total_talk_seconds": (row.total_talk_seconds or 0) if _full_cols else 0,
                "label_name": (row.label_name or "") if _full_cols else "",
                "label_color": (row.label_color or "") if _full_cols else "",
            }

        analyses_map: dict[str, dict] = {
            row.appointment_id: _build_map_row(row)
            for row in ana_result.all()
        }
    else:
        analyses_map = {}


    enriched = []
    for a in all_appts:
        appt_id = str(a["id"])
        parsed = parse_campaign_code(a.get("type", ""))
        # Normalise: strip whitespace so "AVANZ-AVI-0000 " == "AVANZ-AVI-0000"
        campaign_code = parsed.get("raw", "").strip() if parsed.get("valid") else None

        campaign_cfg = _match_campaign_prefix(campaign_code, all_campaigns) if campaign_code else None
        campaign_inactive = (
            campaign_cfg is None
            and bool(_match_campaign_prefix(campaign_code, all_campaigns_inactive))
        ) if campaign_code else False

        if campaign_code and not campaign_cfg:
            logger.warning(
                "No active campaign match | raw_acuity_type=%r | campaign_code=%r | inactive=%s | active_codes=%s",
                a.get("type", ""),
                campaign_code,
                campaign_inactive,
                sorted(all_campaigns.keys()),
            )

        op_display = get_operator_display(a)
        ragione = extract_ragione_sociale(a) or "—"

        raw_labels = a.get("labels") or []
        if raw_labels:
            logger.debug(
                "appt %s labels raw: %s",
                a.get("id"), [(l.get("name"), l.get("color")) for l in raw_labels]
            )
        labels = [
            {"name": lbl.get("name", ""), "color": lbl.get("color") or ""}
            for lbl in raw_labels
        ]

        dt_raw = a.get("datetime", "")
        dt_obj = None
        try:
            dt_obj = parse_iso_datetime(dt_raw)
            is_past = dt_obj < now if dt_obj else False
            dt_display = dt_obj.strftime("%d/%m/%Y %H:%M") if dt_obj else dt_raw[:16].replace("T", " ")
        except Exception:
            is_past = False
            dt_display = dt_raw[:16].replace("T", " ")

        # Acuity "dateCreated" = quando l'appuntamento è stato PRESO (prenotato)
        created_raw = a.get("dateCreated", "")
        created_display = "—"
        _created_date = None   # date object — usato per il raggruppamento

        if created_raw:
            # ── 1. Prova ISO 8601 (formato più comune da Acuity API) ───────────
            # Es: "2024-03-17T14:30:00-0500" oppure "2024-03-17T14:30:00+0000"
            _iso_dt = parse_iso_datetime(created_raw)
            if _iso_dt:
                _created_date = _iso_dt.date()
                created_display = _created_date.strftime("%d/%m/%y")
            else:
                # ── 2. Prova locale italiana, es. "17 marzo 2024" ─────────────
                _IT_MONTHS = {
                    "gennaio":1,"febbraio":2,"marzo":3,"aprile":4,"maggio":5,"giugno":6,
                    "luglio":7,"agosto":8,"settembre":9,"ottobre":10,"novembre":11,"dicembre":12,
                }
                import re as _re
                _m = _re.match(r"(\d{1,2})\s+(\w+)\s+(\d{4})", created_raw.strip(), _re.IGNORECASE)
                if _m:
                    _day2, _mon2, _yr2 = int(_m.group(1)), _m.group(2).lower(), int(_m.group(3))
                    if _mon2 in _IT_MONTHS:
                        from datetime import date as _date2
                        _created_date = _date2(_yr2, _IT_MONTHS[_mon2], _day2)
                        created_display = _created_date.strftime("%d/%m/%y")
                    else:
                        created_display = created_raw[:10]
                else:
                    # ── 3. Altri formati fallback ─────────────────────────────
                    from datetime import datetime as _dt2
                    for _fmt in ("%Y-%m-%d", "%B %d, %Y", "%d/%m/%Y"):
                        try:
                            _created_date = _dt2.strptime(created_raw.strip(), _fmt).date()
                            created_display = _created_date.strftime("%d/%m/%y")
                            break
                        except ValueError:
                            pass
                    else:
                        created_display = created_raw[:10]

        # Date group basato sulla data di PRESA (dateCreated), non data appuntamento
        _today_date = now.date()
        _yest_date = (_today_date - timedelta(days=1))
        try:
            if _created_date is None:
                date_group = "other"
                logger.debug("dateCreated parse failed for appt %s, raw=%r", appt_id, created_raw)
            elif _created_date == _today_date:
                date_group = "today"
            elif _created_date == _yest_date:
                date_group = "yesterday"
            else:
                date_group = "past"
        except Exception:
            date_group = "other"

        # Short date formats for table columns (just day/month, no year)
        _created_short = _created_date.strftime("%d/%m") if _created_date else created_display
        _dt_short = dt_obj.strftime("%d/%m") if dt_obj else (dt_raw[:5] if dt_raw else "")

        enriched.append({
            "id": appt_id,
            "account": a["_account"],
            "dt_display": dt_display,
            "dt_short": _dt_short,
            "created_display": created_display,
            "created_short": _created_short,
            "_created_date_obj": _created_date,   # date object for period filtering
            "campaign_code": campaign_code or a.get("type", "—"),
            "raw_acuity_type": a.get("type", ""),   # exact string from Acuity API
            "campaign_cfg": campaign_cfg,
            "campaign_inactive": campaign_inactive,
            "ragione": ragione,
            "op_display": op_display,
            "labels": labels,
            "is_past": is_past,
            "date_group": date_group,   # future | today | yesterday | past | other
            "phone": extract_phone(a) or "",
            "piva": extract_piva(a) or "",
            "analysis": analyses_map.get(appt_id),
        })

    # ── Period filter (only for v=3 main page) ────────────────────────────
    _period = request.query_params.get("period", "all")
    _date_from_str = request.query_params.get("date_from", "")
    _date_to_str   = request.query_params.get("date_to", "")

    if _period != "all":
        from datetime import date as _date_cls, timedelta as _td
        _today = now.date()
        _yest  = _today - timedelta(days=1)
        _week_start = _today - timedelta(days=_today.weekday())   # Monday
        _month_start = _today.replace(day=1)

        def _in_period(item) -> bool:
            _cd = item.get("_created_date_obj")
            if _cd is None:
                return _period == "all"
            if _period == "today":
                return _cd == _today
            elif _period == "yesterday":
                return _cd == _yest
            elif _period == "week":
                return _cd >= _week_start
            elif _period == "month":
                return _cd >= _month_start
            elif _period == "custom":
                try:
                    _df = _date_cls.fromisoformat(_date_from_str) if _date_from_str else None
                    _dt = _date_cls.fromisoformat(_date_to_str)   if _date_to_str   else None
                except ValueError:
                    return True
                if _df and _cd < _df:
                    return False
                if _dt and _cd > _dt:
                    return False
                return True
            return True

        enriched = [e for e in enriched if _in_period(e)]

    has_processing = any(
        a.get("analysis") and a["analysis"]["status"] == "processing"
        for a in enriched
    )

    # Unique operators and campaign codes for filter dropdowns
    unique_operators = sorted({
        e["op_display"] for e in enriched
        if e["op_display"] and e["op_display"] not in ("—", "")
    })
    unique_campaigns = sorted({
        e["campaign_code"] for e in enriched
        if e["campaign_code"] and e["campaign_code"] not in ("—", "")
    })

    _v = request.query_params.get("v", "1")
    if _v == "3":
        _tmpl = "main_fragment.html"
    elif _v == "2":
        _tmpl = "appointments_v2_fragment.html"
    else:
        _tmpl = "appointments_table_fragment.html"
    return templates.TemplateResponse(
        _tmpl,
        {
            "request": request,
            "appointments": enriched,
            "has_processing": has_processing,
            "unique_operators": unique_operators,
            "unique_campaigns": unique_campaigns,
            "period": _period,
            "label_colors": LABEL_COLORS,
            "label_colors_default": _LABEL_COLORS_DEFAULT,
            "acuity_css_color_map": _ACUITY_CSS_COLOR_MAP,
        },
    )


@router.get("/appointments/status-poll")
async def appointments_status_poll(request: Request, db: AsyncSession = Depends(get_db)):
    """Lightweight JSON endpoint — in-progress + recently completed analyses."""
    if not _is_admin(request):
        return JSONResponse({"items": [], "has_processing": False}, status_code=401)
    from sqlalchemy import or_
    from utils.helpers import utcnow
    import datetime
    cutoff = utcnow() - datetime.timedelta(days=30)
    try:
        result = await db.execute(
            select(Analysis.appointment_id, Analysis.processing_status, Analysis.id,
                   Analysis.qualification_level, Analysis.progress, Analysis.step_message)
            .where(or_(
                Analysis.processing_status.in_(["processing", "pending"]),
                Analysis.updated_at >= cutoff,
            ))
        )
    except Exception as _pe:
        logger.warning("status-poll: updated_at column missing? %s — fallback", _pe)
        await db.rollback()
        result = await db.execute(
            select(Analysis.appointment_id, Analysis.processing_status, Analysis.id,
                   Analysis.qualification_level, Analysis.progress, Analysis.step_message)
            .where(Analysis.processing_status.in_(["processing", "pending"]))
        )
    rows = result.all()
    items = [
        {
            "appt_id": row.appointment_id,
            "status": row.processing_status,
            "analysis_id": row.id,
            "qualification_level": row.qualification_level or "",
            "progress": row.progress or 0,
            "step_message": row.step_message or "",
        }
        for row in rows
    ]
    has_processing = any(r.processing_status in ("processing", "pending") for r in rows)
    return JSONResponse({"items": items, "has_processing": has_processing})


@router.post("/appointments/{account_id}/{appointment_id}/analyze")
async def trigger_appointment_analysis(
    account_id: int,
    appointment_id: str,
    request: Request,
    background_tasks: BackgroundTasks,
    engine: Optional[str] = None,
):
    if not _is_admin(request):
        # AJAX call: return 401 JSON; form POST fallback: redirect
        accept = request.headers.get("accept", "")
        if "application/json" in accept:
            return JSONResponse({"error": "Non autorizzato"}, status_code=401)
        return _login_redirect()

    from services.acuity import get_appointment
    from routers.webhook import run_analysis_pipeline
    from urllib.parse import quote

    full_appointment = await get_appointment(appointment_id, account_id)
    actual_account = account_id
    # Fallback: try the other account in case the appointment is cross-account
    if not full_appointment:
        fallback_account = 2 if account_id == 1 else 1
        full_appointment = await get_appointment(appointment_id, fallback_account)
        if full_appointment:
            actual_account = fallback_account
            logger.info(
                "get_appointment fallback: appt %s found on account %d (tried %d first)",
                appointment_id, fallback_account, account_id,
            )
    if not full_appointment:
        return JSONResponse(
            {"error": f"Impossibile recuperare l'appuntamento {appointment_id} da Acuity."},
            status_code=404,
        )

    # Pre-flight: check campaign exists in DB before starting background task
    from services.campaign_parser import parse_campaign_code
    from services.campaign_db import get_campaign_by_code
    from database import get_db as _get_db

    appt_type = full_appointment.get("type", "")
    campaign_info = parse_campaign_code(appt_type)
    if not campaign_info.get("valid"):
        return JSONResponse(
            {"error": f"Tipo appuntamento non parseable come codice campagna: '{appt_type}'"},
            status_code=422,
        )

    campaign_db = await get_campaign_by_code(campaign_info["raw"])
    if campaign_db is None:
        return JSONResponse(
            {"error": f"Campagna '{campaign_info['raw']}' non configurata nel DB — aggiungila prima di analizzare."},
            status_code=422,
        )

    # engine param: "" or None → no override (use campaign or global default)
    engine_override = engine.strip() if engine and engine.strip() else None

    background_tasks.add_task(
        run_analysis_pipeline,
        appointment_data=full_appointment,
        acuity_account=actual_account,
        engine_override=engine_override,
    )

    return JSONResponse({"ok": True, "appointment_id": appointment_id})


# ── Acuity debug (/admin/ui/acuity-debug) ────────────────────────────────────

@router.get("/acuity-debug", response_class=JSONResponse)
async def acuity_debug(request: Request, appointment_id: str = "", account_id: int = 1):
    """
    GET /admin/ui/acuity-debug?appointment_id=XXXXX&account_id=1
    Mostra tutti i dati raw di un appuntamento Acuity inclusi i form fields.
    """
    if not _is_admin(request):
        return JSONResponse({"error": "Non autorizzato"}, status_code=401)
    if not appointment_id:
        return JSONResponse({"error": "Specificare appointment_id"}, status_code=400)

    from services.acuity import get_appointment, extract_phone, extract_piva, extract_ragione_sociale, find_operator_email, get_operator_display

    data = await get_appointment(appointment_id, account_id)
    if not data:
        return JSONResponse({"error": "Appuntamento non trovato"}, status_code=404)

    # Estrai tutti i form fields in modo leggibile
    all_fields = []
    for form in (data.get("forms") or []):
        form_name = form.get("name") or form.get("id") or "—"
        for val in (form.get("values") or form.get("fields") or []):
            all_fields.append({
                "form": form_name,
                "field_name": val.get("name") or val.get("label") or "—",
                "field_id": val.get("id") or "—",
                "value": val.get("value") or "",
            })
    for val in (data.get("fields") or []):
        all_fields.append({
            "form": "top-level fields",
            "field_name": val.get("name") or val.get("label") or "—",
            "field_id": val.get("id") or "—",
            "value": val.get("value") or "",
        })

    return JSONResponse({
        "appointment_id": data.get("id"),
        "created_at": data.get("createdAt"),
        "datetime": data.get("datetime"),
        "type": data.get("type"),
        "calendar": data.get("calendar"),
        "labels": data.get("labels"),
        "firstName": data.get("firstName"),
        "lastName": data.get("lastName"),
        "email": data.get("email"),
        "phone_top_level": data.get("phone"),
        "extracted_phone": extract_phone(data),
        "extracted_piva": extract_piva(data),
        "extracted_ragione_sociale": extract_ragione_sociale(data),
        "extracted_operator_email": find_operator_email(data),
        "extracted_operator_display": get_operator_display(data),
        "form_fields": all_fields,
        "raw_forms": data.get("forms"),
    })


# ── Sidial test (/admin/ui/sidial-test) ──────────────────────────────────────

@router.get("/sidial-test", response_class=JSONResponse)
async def sidial_test(request: Request, phone: str = ""):
    """
    Testa la connessione Sidial per un numero di telefono.
    GET /admin/ui/sidial-test?phone=0123456789
    Ritorna JSON con leads trovati e registrazioni disponibili.
    """
    if not _is_admin(request):
        return JSONResponse({"error": "non autorizzato"}, status_code=401)

    from services.sidial import (
        _normalize_phone, _search_leads_by_phone, _search_recs_by_lead
    )
    from config import settings as cfg

    result: dict = {
        "sidial_url": cfg.sidial_api_url,
        "token_len": len(cfg.sidial_api_token) if cfg.sidial_api_token else 0,
        "phone_raw": phone,
        "phone_norm": _normalize_phone(phone) if phone else "",
        "leads": [],
        "recordings": [],
        "error": None,
    }

    if not phone:
        result["error"] = "Passa ?phone=NUMERO nella query string"
        return JSONResponse(result)

    try:
        norm = _normalize_phone(phone)
        leads = await _search_leads_by_phone(norm)
        if not leads and norm != phone:
            leads = await _search_leads_by_phone(phone)
        result["leads"] = [
            {"id": l.get("id"), "phone1": l.get("phone1"), "phone2": l.get("phone2"),
             "name": l.get("name") or l.get("companyName") or ""}
            for l in leads
        ]
        for lead in leads:
            lead_id = str(lead.get("id") or "")
            if not lead_id:
                continue
            recs = await _search_recs_by_lead(lead_id)
            for r in recs:
                result["recordings"].append({
                    "id": r.get("id"),
                    "leadId": lead_id,
                    "createdWhen": r.get("createdWhen"),
                    "callLength": r.get("callLength"),
                })
    except Exception as exc:
        result["error"] = str(exc)

    return JSONResponse(result)


# ── Global documents (/admin/ui/global) ──────────────────────────────────────

@router.get("/global", response_class=HTMLResponse)
async def global_docs_list(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(
        select(GlobalDocument).order_by(GlobalDocument.sort_order, GlobalDocument.id)
    )
    docs = result.scalars().all()

    flash_ok  = request.query_params.get("ok", "")
    flash_err = request.query_params.get("err", "")

    return templates.TemplateResponse(
        "global_docs_list.html",
        {
            "request": request,
            "docs": docs,
            "active_page": "global",
            "flash_ok": flash_ok,
            "flash_err": flash_err,
        },
    )


@router.get("/global/new", response_class=HTMLResponse)
async def global_doc_new_form(request: Request):
    if not _is_admin(request):
        return _login_redirect()
    return templates.TemplateResponse(
        "global_docs_form.html",
        {
            "request": request,
            "doc": None,
            "active_page": "global",
            "action": "/admin/ui/global/new",
            "flash_ok": "",
            "flash_err": "",
        },
    )


@router.post("/global/new")
async def global_doc_new_submit(
    request: Request,
    db: AsyncSession = Depends(get_db),
    title: str = Form(""),
    content: str = Form(""),
    sort_order: int = Form(0),
    is_active: str = Form("off"),
):
    if not _is_admin(request):
        return _login_redirect()

    title = title.strip()
    if not title:
        return templates.TemplateResponse(
            "global_docs_form.html",
            {
                "request": request,
                "doc": None,
                "active_page": "global",
                "action": "/admin/ui/global/new",
                "flash_ok": "",
                "flash_err": "Il titolo è obbligatorio.",
            },
            status_code=422,
        )

    doc = GlobalDocument(
        title=title,
        content=content.strip(),
        sort_order=sort_order,
        is_active=(is_active == "on"),
    )
    db.add(doc)
    await db.commit()

    from urllib.parse import quote
    return RedirectResponse(url=f"/admin/ui/global?ok={quote('Documento creato.')}", status_code=303)


@router.get("/global/{doc_id}/edit", response_class=HTMLResponse)
async def global_doc_edit_form(
    request: Request,
    doc_id: int,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(GlobalDocument).where(GlobalDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return RedirectResponse(url="/admin/ui/global", status_code=302)

    return templates.TemplateResponse(
        "global_docs_form.html",
        {
            "request": request,
            "doc": doc,
            "active_page": "global",
            "action": f"/admin/ui/global/{doc_id}/edit",
            "flash_ok": "",
            "flash_err": "",
        },
    )


@router.post("/global/{doc_id}/edit")
async def global_doc_edit_submit(
    request: Request,
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    title: str = Form(""),
    content: str = Form(""),
    sort_order: int = Form(0),
    is_active: str = Form("off"),
):
    if not _is_admin(request):
        return _login_redirect()

    result = await db.execute(select(GlobalDocument).where(GlobalDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return RedirectResponse(url="/admin/ui/global", status_code=302)

    doc.title = title.strip() or doc.title
    doc.content = content.strip()
    doc.sort_order = sort_order
    doc.is_active = (is_active == "on")
    await db.commit()

    from urllib.parse import quote
    return RedirectResponse(url=f"/admin/ui/global?ok={quote('Documento aggiornato.')}", status_code=303)


@router.post("/global/{doc_id}/delete")
async def global_doc_delete(
    request: Request,
    doc_id: int,
    db: AsyncSession = Depends(get_db),
):
    if not _is_admin(request):
        return _login_redirect()

    await db.execute(sa_delete(GlobalDocument).where(GlobalDocument.id == doc_id))
    await db.commit()

    from urllib.parse import quote
    return RedirectResponse(url=f"/admin/ui/global?ok={quote('Documento eliminato.')}", status_code=303)


# ── File text extraction (/admin/ui/upload-extract) ───────────────────────────

@router.post("/upload-extract")
async def upload_extract(
    request: Request,
    file: UploadFile = File(...),
):
    """
    Receive an uploaded file and return its text content as JSON.
    Supported: PDF, DOCX, TXT.  Images and Google Docs not supported.
    Auth checked via cookie (same as other UI routes).
    """
    if not _is_admin(request):
        return JSONResponse({"error": "Non autenticato"}, status_code=401)

    filename = file.filename or ""
    data = await file.read()
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    try:
        if ext == "pdf":
            text = _extract_pdf(data)
        elif ext in ("docx", "doc"):
            text = _extract_docx(data)
        elif ext == "txt":
            text = data.decode("utf-8", errors="ignore")
        else:
            return JSONResponse(
                {"error": f"Formato «{ext}» non supportato. Usa PDF, DOCX o TXT."},
                status_code=400,
            )
    except Exception as exc:
        logger.error("File extraction error (%s): %s", filename, exc)
        return JSONResponse({"error": f"Errore durante l'estrazione: {exc}"}, status_code=500)

    return JSONResponse({"text": text, "filename": filename, "chars": len(text)})


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    """Extract text from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── Private helpers ───────────────────────────────────────────────────────────

def _match_campaign_prefix(campaign_code: str, all_campaigns: dict) -> Optional[Campaign]:
    """
    In-memory longest-prefix match (mirrors campaign_db.get_campaign_by_code).
    Both candidates and dict keys are normalised to UPPER-STRIP so mismatches
    caused by accidental whitespace or lowercase DB entries are avoided.
    """
    if not campaign_code:
        return None
    tokens = [t.strip().upper() for t in campaign_code.strip().split("-") if t.strip()]
    for i in range(len(tokens), 0, -1):
        candidate = "-".join(tokens[:i])
        if candidate in all_campaigns:
            return all_campaigns[candidate]
    return None


def _extract_ragione_sociale(appt: dict) -> str:
    """
    Try to find the company/client name from Acuity form fields,
    falling back to firstName + lastName.
    """
    KEYWORDS = ("ragione", "azienda", "cliente", "societa", "company")
    for form in (appt.get("forms") or []):
        for val in (form.get("values") or []):
            name_lower = re.sub(r"[àèéìòù]", "a", (val.get("name") or "").lower())
            if any(kw in name_lower for kw in KEYWORDS):
                v = (val.get("value") or "").strip()
                if v:
                    return v
    parts = [appt.get("firstName", ""), appt.get("lastName", "")]
    return " ".join(p for p in parts if p).strip() or "—"


# ── Prompt preview (/admin/ui/prompt-preview) ────────────────────────────────

@router.get("/prompt-preview", response_class=HTMLResponse)
async def prompt_preview(
    request: Request,
    db: AsyncSession = Depends(get_db),
    campaign_code: str = "",
):
    """Show the full prompt that would be sent to Claude for a given campaign."""
    if not _is_admin(request):
        return _login_redirect()

    # Load all campaigns for the selector
    camps_result = await db.execute(
        select(Campaign).where(Campaign.code != "_GLOBAL_").order_by(Campaign.code)
    )
    all_campaigns = camps_result.scalars().all()

    prompt_text = None
    selected_campaign = None

    if campaign_code:
        # Find the campaign
        camp_result = await db.execute(
            select(Campaign).where(func.upper(Campaign.code) == campaign_code.strip().upper())
        )
        selected_campaign = camp_result.scalar_one_or_none()

        # Load global docs
        gdocs_result = await db.execute(
            select(GlobalDocument)
            .where(GlobalDocument.is_active == True)
            .order_by(GlobalDocument.sort_order, GlobalDocument.id)
        )
        global_docs = [
            {"title": d.title, "content": d.content}
            for d in gdocs_result.scalars().all()
        ]

        # Load prompt sections
        from services.prompt_db import get_prompt_sections
        try:
            prompt_sections = await get_prompt_sections()
        except Exception:
            prompt_sections = {}

        # Build campaign_info dict (mirrors webhook.py)
        from services.campaign_parser import parse_campaign_code
        campaign_info = parse_campaign_code(campaign_code.strip())

        from services.ai_analysis import build_analysis_prompt
        prompt_text = build_analysis_prompt(
            transcript="[— trascrizione di esempio non disponibile nell'anteprima —]",
            campaign_info=campaign_info,
            script=selected_campaign.script if selected_campaign else None,
            qualification_params=selected_campaign.qualification_params if selected_campaign else None,
            client_info=selected_campaign.client_info if selected_campaign else None,
            operator_email="op.01.mario@effoncall.com",
            prompt_sections=prompt_sections,
            prompt_extra=selected_campaign.prompt_extra if selected_campaign else None,
            global_docs=global_docs,
        )

    return templates.TemplateResponse(
        "prompt_preview.html",
        {
            "request": request,
            "active_page": "prompt",
            "all_campaigns": all_campaigns,
            "campaign_code": campaign_code,
            "selected_campaign": selected_campaign,
            "prompt_text": prompt_text,
            "flash_ok": "",
            "flash_err": "",
        },
    )


def _parse_recipients(raw: str) -> list[str]:
    """Split a textarea (one email per line) into a list of non-empty strings."""
    return [e.strip() for e in raw.splitlines() if e.strip()]


def _form_snapshot(local_vars: dict) -> dict:
    """Collect form field values for re-populating the form on error."""
    keys = ("code", "nome", "script", "qualification_params",
            "client_info", "email_recipients_raw", "notes", "active")
    return {k: local_vars.get(k, "") for k in keys}


# ── Settings page ─────────────────────────────────────────────────────────────

@router.get("/settings", response_class=HTMLResponse)
async def settings_page(request: Request, db: AsyncSession = Depends(get_db)):
    if not _is_admin(request):
        return _login_redirect()

    from models import Setting, Operator as OperatorModel
    from sqlalchemy import text as _text

    # Ensure tables exist (idempotent — runs silently if already present)
    try:
        await db.execute(_text("""
            CREATE TABLE IF NOT EXISTS operators (
                id SERIAL PRIMARY KEY,
                number VARCHAR(10) UNIQUE NOT NULL,
                display_name VARCHAR(100),
                email VARCHAR(200),
                active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMPTZ DEFAULT NOW()
            )
        """))
        await db.execute(_text("""
            CREATE TABLE IF NOT EXISTS settings (
                key VARCHAR(100) PRIMARY KEY,
                value TEXT,
                description TEXT,
                updated_at TIMESTAMPTZ DEFAULT NOW()
            )
        """))
        await db.execute(_text("""
            INSERT INTO settings (key, value, description) VALUES
                ('transcription_engine',     'openai', 'Motore trascrizione default: openai o assemblyai'),
                ('min_call_length_seconds',  '20',     'Durata minima registrazione in secondi'),
                ('sidial_lookback_days',     '90',     'Giorni lookback registrazioni Sidial'),
                ('sidial_retry_count',       '5',      'Numero massimo retry download registrazioni'),
                ('sidial_retry_wait_seconds','180',    'Attesa secondi tra i retry download')
            ON CONFLICT (key) DO NOTHING
        """))
        await db.commit()
    except Exception as _exc:
        logger.warning("settings_page: table creation failed (non-fatal): %s", _exc)
        try:
            await db.rollback()
        except Exception:
            pass

    flash_err = request.query_params.get("err", "")

    try:
        settings_result = await db.execute(select(Setting).order_by(Setting.key))
        all_settings = settings_result.scalars().all()
    except Exception as exc:
        logger.error("settings_page: cannot query settings table: %s", exc)
        all_settings = []
        flash_err = flash_err or f"Tabella settings non disponibile: {exc}"

    try:
        operators_result = await db.execute(
            select(OperatorModel).where(OperatorModel.active == True).order_by(OperatorModel.number)
        )
        operators = operators_result.scalars().all()
    except Exception as exc:
        logger.error("settings_page: cannot query operators table: %s", exc)
        operators = []
        flash_err = flash_err or f"Tabella operators non disponibile: {exc}"

    return templates.TemplateResponse(
        "settings.html",
        {
            "request": request,
            "active_page": "settings",
            "all_settings": all_settings,
            "operators": operators,
            "flash_ok": request.query_params.get("ok", ""),
            "flash_err": flash_err,
        },
    )


@router.post("/settings/update")
async def update_setting_value(
    request: Request,
    key: str = Form(...),
    value: str = Form(...),
):
    if not _is_admin(request):
        return _login_redirect()
    from services.settings_service import set_setting
    await set_setting(key.strip(), value.strip())
    from urllib.parse import quote
    return RedirectResponse(
        url=f"/admin/ui/settings?ok={quote(f'Impostazione {key!r} aggiornata.')}",
        status_code=303,
    )


@router.get("/settings/operators", response_class=JSONResponse)
async def list_operators(request: Request, db: AsyncSession = Depends(get_db)):
    if not _is_admin(request):
        return JSONResponse({"error": "non autorizzato"}, status_code=403)
    from models import Operator as OperatorModel
    result = await db.execute(select(OperatorModel).order_by(OperatorModel.number))
    ops = result.scalars().all()
    return JSONResponse([
        {
            "id": o.id, "number": o.number,
            "display_name": o.display_name,
            "email": o.email,
            "active": o.active,
        }
        for o in ops
    ])


@router.post("/settings/operators/save")
async def save_operator(
    request: Request,
    db: AsyncSession = Depends(get_db),
    number: str = Form(...),
    display_name: str = Form(""),
    email: str = Form(""),
):
    if not _is_admin(request):
        return _login_redirect()
    from models import Operator as OperatorModel
    number = number.strip()
    existing = await db.execute(
        select(OperatorModel).where(OperatorModel.number == number)
    )
    op = existing.scalar_one_or_none()
    if op:
        op.display_name = display_name.strip() or None
        op.email = email.strip() or None
        op.active = True
    else:
        db.add(OperatorModel(
            number=number,
            display_name=display_name.strip() or None,
            email=email.strip() or None,
        ))
    await db.commit()
    from urllib.parse import quote
    return RedirectResponse(
        url=f"/admin/ui/settings?ok={quote(f'Operatore #{number} salvato.')}",
        status_code=303,
    )


# ── Pipeline steps endpoint ───────────────────────────────────────────────────

@router.get("/analyses/{analysis_id}/pipeline-steps", response_class=JSONResponse)
async def get_pipeline_steps(
    analysis_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Return pipeline_steps JSON for an analysis (for polling UI)."""
    if not _is_admin(request):
        return JSONResponse({"error": "non autorizzato"}, status_code=403)
    result = await db.execute(select(Analysis).where(Analysis.id == analysis_id))
    analysis = result.scalar_one_or_none()
    if not analysis:
        return JSONResponse({"error": "analisi non trovata"}, status_code=404)
    return JSONResponse({
        "id": analysis_id,
        "processing_status": analysis.processing_status,
        "qualification_level": analysis.qualification_level or "",
        "pipeline_steps": analysis.pipeline_steps or {},
        "progress": analysis.progress or 0,
        "step_message": analysis.step_message or "",
    })


# ── Main page (unified view) ──────────────────────────────────────────────────

@router.get("/main", response_class=HTMLResponse)
async def main_page(request: Request):
    if not _is_admin(request):
        return _login_redirect()
    return templates.TemplateResponse(
        "main.html",
        {"request": request, "active_page": "main"},
    )
